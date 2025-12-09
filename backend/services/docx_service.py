"""Service for generating DOCX reports from analysis data."""
import io
import asyncio
from models.schemas import AnalysisData

# Store for DOCX files (in production, use Redis or database)
docx_cache = {}

def generate_docx_bytes(analysis_data: AnalysisData) -> io.BytesIO:
    """Generate DOCX file in memory."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Interview Analysis Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 1. Executive Summary
    if hasattr(analysis_data, 'executiveSummary') and analysis_data.executiveSummary:
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(analysis_data.executiveSummary)

    # 2. Statistics
    if analysis_data.statistics:
        doc.add_heading('Expert Statistics', 1)
        stats = analysis_data.statistics
        
        # Metrics at the top (Handle dict or object access)
        def get_stat(key):
            return stats.get(key, 'N/A') if isinstance(stats, dict) else getattr(stats, key, 'N/A')

        p = doc.add_paragraph()
        p.add_run("Communication Score: ").bold = True
        p.add_run(f"{get_stat('communicationScore')}/100")
        
        p = doc.add_paragraph()
        p.add_run("Technical Depth: ").bold = True
        p.add_run(f"{get_stat('technicalDepth')}/100")
        
        p = doc.add_paragraph()
        p.add_run("Engagement Score: ").bold = True
        p.add_run(f"{get_stat('engagementScore')}/100")

    # 3. Key Strengths & Growth Areas
    if hasattr(analysis_data, 'strengthsWeaknesses') and analysis_data.strengthsWeaknesses:
        sw = analysis_data.strengthsWeaknesses
        doc.add_heading('Key Strengths & Growth Areas', 1)
        
        doc.add_heading('Strengths', 2)
        for s in (sw.get('strengths', []) if isinstance(sw, dict) else sw.strengths):
            doc.add_paragraph(s, style='List Bullet')
            
        doc.add_heading('Growth Areas', 2)
        for w in (sw.get('weaknesses', []) if isinstance(sw, dict) else sw.weaknesses):
            doc.add_paragraph(w, style='List Bullet')

    # 4. Thinking Process Analysis
    if hasattr(analysis_data, 'thinkingProcess') and analysis_data.thinkingProcess:
        tp = analysis_data.thinkingProcess
        doc.add_heading('Thinking Process Analysis', 1)
        
        # Access helpers
        def get_tp(key):
            return tp.get(key, 'N/A') if isinstance(tp, dict) else getattr(tp, key, 'N/A')

        doc.add_paragraph(f"Methodology: {get_tp('methodology')}")
        doc.add_paragraph(f"Logical Structure: {get_tp('structure')}")
        
        # Clean Edge Case logic
        edge = get_tp('edgeCaseExplanation')
        if "not explicitly tested" in str(edge).lower():
            edge = "N/A"
        doc.add_paragraph(f"Edge Cases: {edge}")

    # 5. General Assessment
    if analysis_data.generalComments:
        doc.add_heading('General Assessment', 1)
        # Handle dict or object access
        gc = analysis_data.generalComments
        def get_gc(key):
            return gc.get(key, 'N/A') if isinstance(gc, dict) else getattr(gc, key, 'N/A')
            
        doc.add_paragraph(f"Interview Overview: {get_gc('howInterview')}")
        doc.add_paragraph(f"Interviewer's Attitude: {get_gc('attitude')}")
        doc.add_paragraph(f"Structure: {get_gc('structure')}")
        doc.add_paragraph(f"Platform: {get_gc('platform')}")

    # 6. Key Technical Emphasis Points
    if analysis_data.keyPoints:
        doc.add_heading('Key Technical Emphasis Points', 1)
        for point in analysis_data.keyPoints:
            p = doc.add_paragraph(style='List Bullet')
            title = point.get('title', 'Point') if isinstance(point, dict) else point.title
            content = point.get('content', '') if isinstance(point, dict) else point.content
            p.add_run(f"{title}: ").bold = True
            p.add_run(content)
    
    # 7. Technologies
    if analysis_data.technologies:
        doc.add_heading('Technologies and Tools Used', 1)
        for tech in analysis_data.technologies:
            name = tech.get('name', 'Unknown') if isinstance(tech, dict) else tech.name
            timestamps = tech.get('timestamps', '') if isinstance(tech, dict) else tech.timestamps
            
            tech_text = f"{name}"
            if timestamps:
                tech_text += f" ({timestamps})"
            doc.add_paragraph(tech_text, style='List Bullet')
    
    # 8. Coding Challenge
    if analysis_data.codingChallenge:
        doc.add_heading('Live Coding Challenge Details', 1)
        doc.add_paragraph(f"Core Exercise: {analysis_data.codingChallenge.get('coreExercise', 'N/A')}")
        doc.add_paragraph(f"Critical Follow-up: {analysis_data.codingChallenge.get('followUp', 'N/A')}")
        doc.add_paragraph(f"Required Knowledge: {analysis_data.codingChallenge.get('knowledge', 'N/A')}")

    # 9. Q&A Topics
    if analysis_data.qaTopics:
        doc.add_heading('Non-Technical & Situational Q&A Topics', 1)
        for topic in analysis_data.qaTopics:
            p = doc.add_paragraph(style='List Bullet')
            title = topic.get('title', 'Topic') if isinstance(topic, dict) else topic.title
            content = topic.get('content', '') if isinstance(topic, dict) else topic.content
            p.add_run(f"{title}: ").bold = True
            p.add_run(content)
    
    # Save to BytesIO
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


async def generate_docx_async(analysis_data: AnalysisData, full_transcript: str, cache_key: str):
    """Generate DOCX file asynchronously in background."""
    try:
        print(f"📝 Generating DOCX in background for key: {cache_key}")
        
        # Run generation in thread pool to avoid blocking
        docx_buffer = await asyncio.to_thread(generate_docx_bytes, analysis_data)
        
        # Cache the DOCX
        docx_cache[cache_key] = docx_buffer.getvalue()
        print(f"✅ DOCX generated and cached for key: {cache_key}")
        
    except Exception as e:
        print(f"❌ Error generating DOCX: {str(e)}")


def get_cached_docx(cache_key: str) -> bytes | None:
    """Retrieve a cached DOCX file."""
    return docx_cache.get(cache_key)
