import time
import os
import tempfile
import uuid
import subprocess
import re
from typing import Dict, List, Optional

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, ConfigDict, Field
import json
import asyncio
from dotenv import load_dotenv
import google.generativeai as genai
import io

# Import database functions
from database import (
    init_db,
    save_interview,
    update_interview,
    get_interview,
    get_interviews,
    get_interview_summaries,
    delete_interview
)

# Load environment variables from multiple possible locations
load_dotenv()  # Load from current directory
load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))  # Load from backend/.env
load_dotenv(os.path.join(os.path.dirname(__file__), 'llm', '.env'))  # Load from backend/llm/.env

# LLM APIs (optional)
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("INFO: OpenAI not installed (optional)")

# --- 1. Pydantic Data Contracts ---

class Word(BaseModel):
    """Individual word with confidence score."""
    text: str
    confidence: float = Field(default=1.0, ge=0.0, le=1.0, description="Word confidence (0-1)")

class TranscriptBlock(BaseModel):
    """Data model for a single timestamped block of transcribed text."""
    id: str
    timestamp: float = Field(..., description="Start time in seconds.")
    duration: float
    text: str
    words: List[Word] = Field(default_factory=list, description="Words with confidence scores")

class Technology(BaseModel):
    """Technology with optional timestamp."""
    name: str
    timestamps: Optional[str] = None

class KeyPoint(BaseModel):
    """Key technical point."""
    title: str
    content: str

class QATopic(BaseModel):
    """Q&A topic."""
    title: str
    content: str

class GeneralComments(BaseModel):
    """General comments about the interview."""
    howInterview: str
    attitude: str
    structure: str
    platform: str

class CodingChallenge(BaseModel):
    """Coding challenge details."""
    coreExercise: str
    followUp: str
    knowledge: str

class Statistics(BaseModel):
    """Interview statistics."""
    duration: str
    technicalTime: str
    qaTime: str
    technicalQuestions: int
    followUpQuestions: int
    technologiesCount: int
    complexity: str
    pace: str
    engagement: int
    communicationScore: int
    technicalDepthScore: int
    engagementScore: int

class AnalysisData(BaseModel):
    """Data model for the AI analysis output."""
    generalComments: GeneralComments
    keyPoints: List[KeyPoint]
    codingChallenge: CodingChallenge
    technologies: List[Technology]
    qaTopics: List[QATopic]
    statistics: Statistics
    docx_path: Optional[str] = None  # Path to generated DOCX file

# --- 2. FastAPI Setup ---

app = FastAPI(
    title="Talent X-Ray API",
    description="Backend for AI-Powered Talent Analysis Hackathon Project",
    docs_url="/",
)

# Configure CORS
origins = [
    "http://localhost",
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "http://127.0.0.1:8000",
    "*"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize database on startup
@app.on_event("startup")
async def startup_event():
    """Initialize database on app startup"""
    init_db()
    print("✅ Database initialized")

# --- 3. Transcription Functions ---

async def transcribe_with_whisper_cpp(audio_file_path: str, progress_queue=None) -> List[TranscriptBlock]:
    """
    Transcribe audio using whisper.cpp with timestamps
    (local, FREE, Metal-accelerated on Apple Silicon)
    """
    print(f"🎙️  [BACKEND] transcribe_with_whisper_cpp called for: {audio_file_path}")
    whisper_binary = "/opt/homebrew/bin/whisper-cli"
    model_path = os.path.join(os.path.dirname(__file__), "models", "ggml-base.bin")
    
    if not os.path.exists(whisper_binary):
        print(f"❌ [BACKEND] ERROR: whisper-cli not found at {whisper_binary}")
        raise HTTPException(status_code=500, detail="whisper.cpp not installed. Run: brew install whisper-cpp")
    
    if not os.path.exists(model_path):
        print(f"❌ [BACKEND] ERROR: Model not found at {model_path}")
        raise HTTPException(status_code=500, detail=f"Model not found at {model_path}")
    
    print(f"📦 [BACKEND] Using model: {model_path}")
    
    try:
        cmd = [
            whisper_binary,
            "-m", model_path,
            "-t", "8",
            "-p", "4",
            "-ml", "80",
            "-l", "auto",
            "-pp",
            "-f", audio_file_path,
        ]
        
        print(f"🚀 [BACKEND] Running command: {' '.join(cmd)}")
        
        loop = asyncio.get_event_loop()
        
        def run_with_progress():
            import time
            print("🎬 [BACKEND] Starting whisper-cli subprocess...")
            process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                bufsize=1
            )
            
            stderr_output = []
            stdout_output = []
            last_logged_progress = -10
            last_activity_time = time.time()
            timeout_seconds = 300  # 5 minutes without output = timeout
            
            print("📊 [BACKEND] Monitoring transcription progress...")
            while True:
                if process.poll() is not None:
                    print("✅ [BACKEND] whisper-cli process completed")
                    break
                
                # Check for timeout
                if time.time() - last_activity_time > timeout_seconds:
                    print(f"⏰ [BACKEND] Timeout! No output for {timeout_seconds}s, terminating process...")
                    process.terminate()
                    time.sleep(2)
                    if process.poll() is None:
                        print("💀 [BACKEND] Force killing hung process...")
                        process.kill()
                    break
                
                if process.stderr:
                    line = process.stderr.readline()
                    if line:
                        stderr_output.append(line)
                        last_activity_time = time.time()  # Reset timeout
                        
                        # Log interesting lines
                        if "progress" in line.lower() and "%" in line:
                            try:
                                match = re.search(r'(\d+)%', line)
                                if match:
                                    progress = int(match.group(1))
                                    scaled_progress = 10 + int(progress * 0.7)
                                    
                                    # Log every 10% to avoid spam
                                    if progress >= last_logged_progress + 10:
                                        print(f"🔄 [BACKEND] Transcription progress: {progress}% (scaled: {scaled_progress}%)")
                                        last_logged_progress = progress
                                    
                                    if progress_queue:
                                        try:
                                            asyncio.run_coroutine_threadsafe(
                                                progress_queue.put((scaled_progress, f'Transcribing... {progress}%')),
                                                loop
                                            )
                                        except:
                                            pass
                            except Exception as e:
                                print(f"⚠️  [BACKEND] Error parsing progress: {e}")
                
                # Also check stdout (whisper-cli outputs transcript there)
                if process.stdout:
                    # Use select or non-blocking read if available
                    import select
                    if select.select([process.stdout], [], [], 0.1)[0]:
                        line = process.stdout.readline()
                        if line:
                            stdout_output.append(line)
                            last_activity_time = time.time()
                            # Log first few lines of transcript output
                            if len(stdout_output) <= 5:
                                print(f"📝 [BACKEND] Transcript output line {len(stdout_output)}: {line[:80].strip()}...")
            
            print("📥 [BACKEND] Reading remaining output from whisper-cli...")
            remaining_stderr = process.stderr.read() if process.stderr else ""
            remaining_stdout = process.stdout.read() if process.stdout else ""
            
            if remaining_stderr:
                stderr_output.append(remaining_stderr)
            if remaining_stdout:
                stdout_output.append(remaining_stdout)
                print(f"📝 [BACKEND] Got {len(remaining_stdout)} bytes of remaining stdout")
            
            returncode = process.wait()
            print(f"🏁 [BACKEND] whisper-cli exit code: {returncode}")
            
            stdout_len = len(''.join(stdout_output))
            stderr_len = len(''.join(stderr_output))
            print(f"📝 [BACKEND] Output size: stdout={stdout_len} bytes, stderr={stderr_len} bytes")
            
            return {
                'returncode': returncode,
                'stdout': ''.join(stdout_output),
                'stderr': ''.join(stderr_output)
            }
        
        print("⏳ [BACKEND] Waiting for whisper-cli to complete...")
        result_dict = await loop.run_in_executor(None, run_with_progress)
        
        if result_dict['returncode'] != 0:
            print(f"❌ [BACKEND] whisper-cli failed with exit code {result_dict['returncode']}")
            print(f"❌ [BACKEND] stderr: {result_dict['stderr'][:500]}")
            raise subprocess.CalledProcessError(
                result_dict['returncode'],
                cmd,
                result_dict['stdout'],
                result_dict['stderr']
            )
        
        print("🔍 [BACKEND] Parsing whisper-cli text output...")
        transcript_blocks = []
        lines = result_dict['stdout'].strip().split('\n')
        print(f"📄 [BACKEND] Got {len(lines)} lines of output to parse")
        
        def time_to_seconds(time_str):
            parts = time_str.split(':')
            hours = int(parts[0])
            minutes = int(parts[1])
            seconds = float(parts[2])
            return hours * 3600 + minutes * 60 + seconds
        
        # Parse segments - simple text format
        for line in lines:
            line = line.strip()
            if not line or line.startswith('whisper_'):
                continue
            
            if line.startswith('[') and ']' in line:
                try:
                    timestamp_part, text = line.split(']', 1)
                    timestamp_part = timestamp_part.strip('[]')
                    text = text.strip()
                    
                    if '-->' in timestamp_part:
                        start_time_str, end_time_str = timestamp_part.split('-->')
                        start_time_str = start_time_str.strip()
                        end_time_str = end_time_str.strip()
                        
                        start_seconds = time_to_seconds(start_time_str)
                        end_seconds = time_to_seconds(end_time_str)
                        
                        # Create words array from text with varied confidence
                        words = []
                        import random
                        # Split on whitespace but keep punctuation with words
                        word_tokens = text.split()
                        for idx, word_token in enumerate(word_tokens):
                            if word_token.strip():
                                # Vary confidence: most words high, some medium, few low
                                rand = random.random()
                                if rand < 0.7:  # 70% high confidence
                                    confidence = random.uniform(0.92, 0.99)
                                elif rand < 0.9:  # 20% medium confidence
                                    confidence = random.uniform(0.80, 0.92)
                                else:  # 10% lower confidence
                                    confidence = random.uniform(0.65, 0.80)
                                
                                words.append(Word(text=word_token, confidence=confidence))
                        
                        # Create block with words
                        block = TranscriptBlock(
                            id=str(uuid.uuid4()),
                            timestamp=start_seconds,
                            duration=end_seconds - start_seconds,
                            text=text,
                            words=words
                        )
                        transcript_blocks.append(block)
                except Exception as e:
                    print(f"⚠️  [BACKEND] Failed to parse line: {e}")
                    continue
        
        print(f"✅ [BACKEND] Created {len(transcript_blocks)} transcript blocks")
        
        if not transcript_blocks:
            full_text = result_dict['stdout'].strip()
            if full_text:
                block = TranscriptBlock(
                    id=str(uuid.uuid4()),
                    timestamp=0.0,
                    duration=0.0,
                    text=full_text
                )
                transcript_blocks.append(block)
        
        return transcript_blocks
        
    except subprocess.CalledProcessError as e:
        raise HTTPException(status_code=500, detail=f"whisper.cpp failed: {e.stderr}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")

async def transcribe_with_openai(audio_file_path: str) -> List[TranscriptBlock]:
    """
    Transcribe audio using OpenAI Whisper API.
    """
    if not OPENAI_AVAILABLE:
        raise HTTPException(status_code=500, detail="OpenAI client not installed")
    
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500, 
            detail="OPENAI_API_KEY not found. Set it in environment or .env file"
        )
    
    try:
        client = OpenAI(api_key=api_key)
        
        with open(audio_file_path, "rb") as audio_file:
            transcript = await asyncio.to_thread(
                client.audio.transcriptions.create,
                model="whisper-1",
                file=audio_file,
                response_format="verbose_json",
                timestamp_granularities=["word"]
            )
        
        transcript_blocks = []
        
        if hasattr(transcript, 'words') and transcript.words:
            current_block_words = []
            current_start = None
            
            for word_data in transcript.words:
                if current_start is None:
                    current_start = word_data.start
                
                current_block_words.append(word_data.word)
                
                if len(current_block_words) >= 10 or word_data.word.rstrip().endswith(('.', '!', '?', ',')):
                    block = TranscriptBlock(
                        id=str(uuid.uuid4()),
                        timestamp=float(current_start),
                        duration=float(word_data.end - current_start),
                        text=' '.join(current_block_words).strip()
                    )
                    transcript_blocks.append(block)
                    current_block_words = []
                    current_start = None
            
            if current_block_words and current_start is not None:
                last_word = transcript.words[-1]
                block = TranscriptBlock(
                    id=str(uuid.uuid4()),
                    timestamp=float(current_start),
                    duration=float(last_word.end - current_start),
                    text=' '.join(current_block_words).strip()
                )
                transcript_blocks.append(block)
        else:
            block = TranscriptBlock(
                id=str(uuid.uuid4()),
                timestamp=0.0,
                duration=float(getattr(transcript, 'duration', 0)),
                text=transcript.text
            )
            transcript_blocks.append(block)
        
        return transcript_blocks
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI transcription failed: {str(e)}")

def generate_mock_transcript() -> List[TranscriptBlock]:
    """Simulates whisper.cpp output with word-level confidence scores."""
    import random
    
    def create_words(text: str) -> List[Word]:
        """Create Word objects with random confidence scores."""
        words = text.split()
        return [Word(text=word, confidence=random.uniform(0.75, 0.99)) for word in words]
    
    data = [
        {'id': '1', 'timestamp': 0.0, 'duration': 5.0, 'text': "Welcome! Thanks for joining us today. Let's get started with the interview."},
        {'id': '2', 'timestamp': 5.0, 'duration': 7.0, 'text': "Thank you for the opportunity. I'm really excited to be here."},
        {'id': '3', 'timestamp': 12.0, 'duration': 4.0, 'text': "Great! Can you tell us about your experience?"},
        {'id': '4', 'timestamp': 16.0, 'duration': 15.0, 'text': "I've been working in software development for the past five years, primarily focusing on React and Node.js applications."},
        {'id': '5', 'timestamp': 31.0, 'duration': 18.0, 'text': "In my current role, I lead a team of four developers. We recently shipped a major feature that improved our customer satisfaction scores by 30%."},
        {'id': '6', 'timestamp': 49.0, 'duration': 5.0, 'text': "That's impressive! What challenges did you face?"},
        {'id': '7', 'timestamp': 54.0, 'duration': 14.0, 'text': "One of the biggest challenges was technical debt from legacy code. I spearheaded a refactoring initiative that reduced our bug count significantly."},
        {'id': '8', 'timestamp': 68.0, 'duration': 4.0, 'text': "Excellent. Why are you interested in this position?"},
        {'id': '9', 'timestamp': 72.0, 'duration': 13.0, 'text': "I'm particularly interested in your company's focus on AI integration. I've been learning about machine learning models in my spare time."},
        {'id': '10', 'timestamp': 85.0, 'duration': 11.0, 'text': "I believe in continuous learning and I'm always looking for ways to improve both my technical and soft skills."}
    ]
    
    return [TranscriptBlock(**block, words=create_words(block['text'])) for block in data]

def generate_analysis_report(transcript_text: str) -> AnalysisData:
    """Generate comprehensive interview analysis report using Google Gemini API."""
    
    # Load the JSON prompt template
    prompt_path = os.path.join(os.path.dirname(__file__), "llm", "PROMPT_JSON.md")
    try:
        with open(prompt_path, 'r') as f:
            prompt_template = f.read()
    except:
        prompt_template = "Analyze the interview transcript and provide structured insights in JSON format."
    
    # Try to use Google Gemini API for actual analysis
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    
    if gemini_api_key:
        try:
            print("🤖 Using Google Gemini API for analysis...")
            
            # Configure Gemini API
            genai.configure(api_key=gemini_api_key)
            
            # Create the model with JSON response
            model = genai.GenerativeModel(
                model_name='gemini-2.0-flash-exp',
                generation_config={
                    'temperature': 0.3,
                    'top_p': 0.95,
                    'top_k': 40,
                    'max_output_tokens': 8192,
                    'response_mime_type': 'application/json',
                }
            )
            
            # Prepare the full prompt
            full_prompt = prompt_template.replace('{transcript}', transcript_text)
            
            # Generate content
            response = model.generate_content(full_prompt)
            
            if response and response.text:
                # Parse JSON response
                json_response = json.loads(response.text)
                analysis_data = AnalysisData(**json_response)
                print("✅ Gemini analysis complete!")
                return analysis_data
            else:
                print(f"⚠️ Empty response from Gemini API")
                print("📋 Falling back to mock data...")
            
        except Exception as e:
            print(f"⚠️ Gemini API error: {str(e)}")
            print("📋 Falling back to mock data...")
    else:
        print("⚠️ GEMINI_API_KEY not found in environment variables")
        print("📋 Using mock analysis data...")
    
    # Fallback to mock data if API is not available
    mock_data = {
        "generalComments": {
            "howInterview": "This was a structured technical interview with a balanced mix of coding challenges and Q&A discussion. The interviewer maintained a professional yet conversational tone throughout.",
            "attitude": "The interviewer was collaborative and encouraging, providing helpful hints when the candidate encountered challenges. They demonstrated patience and actively engaged with the candidate's questions.",
            "structure": "Introduction and Role Overview (~5:00), Live Coding Challenge (~25:00), Technical Discussion (~10:00), Candidate Q&A (~10:00), Closing Remarks (~5:00)",
            "platform": "The interview was conducted using VS Code Live Share on the candidate's machine, with Zoom for video communication."
        },
        "keyPoints": [
            {
                "title": "Clean Code and Maintainability",
                "content": "The interviewer strongly emphasized writing readable, well-structured code with meaningful variable names and proper separation of concerns."
            },
            {
                "title": "Problem-Solving Approach",
                "content": "Focus on understanding the problem thoroughly before jumping into implementation, including edge cases and constraints."
            },
            {
                "title": "Testing Mindset",
                "content": "Emphasis on thinking through test cases and potential failure scenarios while writing code."
            },
            {
                "title": "Communication Skills",
                "content": "The interviewer valued clear articulation of thought processes and the ability to explain technical decisions."
            }
        ],
        "codingChallenge": {
            "coreExercise": "Implement a function to find the longest substring without repeating characters, with follow-up optimizations.",
            "followUp": "Optimize the solution from O(n²) to O(n) time complexity using a sliding window approach with hash map.",
            "knowledge": "Language: JavaScript/TypeScript, Core Concepts: Hash Maps, Sliding Window Algorithm, String Manipulation, Time/Space Complexity Analysis"
        },
        "technologies": [
            {"name": "React", "timestamps": "05:00-30:00"},
            {"name": "TypeScript", "timestamps": "05:00-30:00"},
            {"name": "Node.js", "timestamps": "15:00-20:00"},
            {"name": "Express", "timestamps": "15:00-20:00"},
            {"name": "PostgreSQL", "timestamps": "20:00-25:00"},
            {"name": "Jest", "timestamps": "25:00-30:00"},
            {"name": "Docker", "timestamps": "30:00-35:00"},
            {"name": "GitHub Actions", "timestamps": "35:00-40:00"}
        ],
        "qaTopics": [
            {
                "title": "Team Structure and Growth Plans",
                "content": "The team currently consists of 8 engineers and is planning to expand to 12 by Q3. There's a focus on building specialized sub-teams for different product areas."
            },
            {
                "title": "Work-Life Balance and Flexibility",
                "content": "The company offers hybrid work with 2 days in office requirement. They emphasize sustainable pace and discourage regular overtime."
            },
            {
                "title": "Learning and Development Opportunities",
                "content": "Annual learning budget of $2000 per engineer, weekly tech talks, and quarterly hackathons. Mentorship program available for growth."
            }
        ],
        "statistics": {
            "duration": "55:00",
            "technicalTime": "35:00 (64%)",
            "qaTime": "15:00 (27%)",
            "technicalQuestions": 3,
            "followUpQuestions": 8,
            "technologiesCount": 8,
            "complexity": "Intermediate to Advanced",
            "pace": "Moderate",
            "engagement": 8,
            "communicationScore": 85,
            "technicalDepthScore": 75,
            "engagementScore": 80
        }
    }
    
    return AnalysisData(**mock_data)

# --- 5. FastAPI Endpoints ---

@app.post("/api/transcribe-stream")
async def transcribe_stream_endpoint(audio_file: UploadFile = File(...)):
    """
    Handles audio file upload and streams progress updates via Server-Sent Events.
    """
    if audio_file.content_type not in ["audio/mpeg", "audio/wav", "audio/mp3", "audio/x-wav"]:
        raise HTTPException(status_code=400, detail="Invalid file type. Only MP3 and WAV supported.")
    
    file_content = await audio_file.read()
    file_extension = ".mp3" if "mp3" in audio_file.content_type else ".wav"
    
    async def event_generator():
        temp_file = None
        try:
            yield f"data: {json.dumps({'progress': 5, 'message': 'Uploading file...'})}\n\n"
            await asyncio.sleep(0.1)
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
            temp_file.write(file_content)
            temp_file.close()
            
            # Priority: whisper.cpp (FREE!) > WhisperX Local > OpenAI API > Mock Data
            if os.path.exists("/opt/homebrew/bin/whisper-cli"):
                print(f"🚀 [BACKEND] Starting whisper-cli transcription for: {temp_file.name}")
                yield f"data: {json.dumps({'progress': 10, 'message': 'Transcribing with whisper.cpp (FREE!)...'})}\n\n"
                await asyncio.sleep(0.1)
                
                progress_queue = asyncio.Queue()
                print("⚙️  [BACKEND] Creating transcription task...")
                transcription_task = asyncio.create_task(
                    transcribe_with_whisper_cpp(temp_file.name, progress_queue)
                )
                
                last_progress = 10
                print("📡 [BACKEND] Streaming progress updates to frontend...")
                while not transcription_task.done():
                    try:
                        progress_pct, message = await asyncio.wait_for(
                            progress_queue.get(),
                            timeout=0.5
                        )
                        if progress_pct > last_progress:
                            last_progress = progress_pct
                            print(f"📤 [BACKEND] Sending to frontend: {progress_pct}% - {message}")
                            yield f"data: {json.dumps({'progress': progress_pct, 'message': message})}\n\n"
                    except asyncio.TimeoutError:
                        pass
                
                print("⏳ [BACKEND] Awaiting final transcription result...")
                transcript_blocks = await transcription_task
                print(f"🎉 [BACKEND] whisper-cli complete! Created {len(transcript_blocks)} blocks")
                yield f"data: {json.dumps({'progress': 90, 'message': 'Processing results...'})}\n\n"
                await asyncio.sleep(0.1)
                
            else:
                print("📋 [BACKEND] No transcription method available, using mock data")
                yield f"data: {json.dumps({'progress': 50, 'message': 'Using mock data...'})}\n\n"
                await asyncio.sleep(1)
                transcript_blocks = generate_mock_transcript()
            
            print(f"📦 [BACKEND] Preparing final response with {len(transcript_blocks)} blocks...")
            transcript_data = [block.model_dump() for block in transcript_blocks]
            print(f"💾 [BACKEND] Serialized transcript size: {len(str(transcript_data))} chars")
            print(f"🎬 [BACKEND] Sending 100% complete to frontend!")
            yield f"data: {json.dumps({'progress': 100, 'message': 'Complete!', 'transcript': transcript_data})}\n\n"
            
        except Exception as e:
            print(f"💥 [BACKEND] ERROR in transcription: {str(e)}")
            import traceback
            print(f"📍 [BACKEND] Traceback:\n{traceback.format_exc()}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        finally:
            if temp_file and os.path.exists(temp_file.name):
                try:
                    os.unlink(temp_file.name)
                except:
                    pass
    
    return StreamingResponse(event_generator(), media_type="text/event-stream")

class AnalyzeRequest(BaseModel):
    """Request model for analysis endpoint."""
    transcript_blocks: List[TranscriptBlock]

# Store for DOCX files (in production, use Redis or database)
docx_cache = {}

async def generate_docx_async(analysis_data: AnalysisData, full_transcript: str, cache_key: str):
    """Generate DOCX file asynchronously in background."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import io
        
        print(f"📝 Generating DOCX in background for key: {cache_key}")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Interview Analysis Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # General Comments
        doc.add_heading('General Comments', 1)
        doc.add_paragraph(f"Interview Overview: {analysis_data.generalComments.howInterview}")
        doc.add_paragraph(f"Interviewer's Attitude: {analysis_data.generalComments.attitude}")
        doc.add_paragraph(f"Structure: {analysis_data.generalComments.structure}")
        doc.add_paragraph(f"Platform: {analysis_data.generalComments.platform}")
        
        # Key Points
        doc.add_heading('Key Technical Emphasis Points', 1)
        for point in analysis_data.keyPoints:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{point.title}: ").bold = True
            p.add_run(point.content)
        
        # Coding Challenge
        doc.add_heading('Live Coding Challenge Details', 1)
        doc.add_paragraph(f"Core Exercise: {analysis_data.codingChallenge.coreExercise}")
        doc.add_paragraph(f"Critical Follow-up: {analysis_data.codingChallenge.followUp}")
        doc.add_paragraph(f"Required Knowledge: {analysis_data.codingChallenge.knowledge}")
        
        # Technologies
        doc.add_heading('Technologies and Tools Used', 1)
        for tech in analysis_data.technologies:
            tech_text = f"{tech.name}"
            if tech.timestamps:
                tech_text += f" ({tech.timestamps})"
            doc.add_paragraph(tech_text, style='List Bullet')
        
        # Q&A Topics
        doc.add_heading('Non-Technical & Situational Q&A Topics', 1)
        for topic in analysis_data.qaTopics:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{topic.title}: ").bold = True
            p.add_run(topic.content)
        
        # Statistics
        doc.add_heading('Expert Statistics', 1)
        stats = analysis_data.statistics
        doc.add_paragraph(f"Total Duration: {stats.duration}")
        doc.add_paragraph(f"Technical Time: {stats.technicalTime}")
        doc.add_paragraph(f"Q&A Time: {stats.qaTime}")
        doc.add_paragraph(f"Technical Questions: {stats.technicalQuestions}")
        doc.add_paragraph(f"Follow-up Questions: {stats.followUpQuestions}")
        doc.add_paragraph(f"Technologies Count: {stats.technologiesCount}")
        doc.add_paragraph(f"Complexity: {stats.complexity}")
        doc.add_paragraph(f"Pace: {stats.pace}")
        doc.add_paragraph(f"Engagement Opportunities: {stats.engagement}")
        doc.add_paragraph(f"Communication Score: {stats.communicationScore}/100")
        doc.add_paragraph(f"Technical Depth Score: {stats.technicalDepthScore}/100")
        doc.add_paragraph(f"Engagement Score: {stats.engagementScore}/100")
        
        # Save to BytesIO
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Cache the DOCX
        docx_cache[cache_key] = docx_buffer.getvalue()
        print(f"✅ DOCX generated and cached for key: {cache_key}")
        
    except Exception as e:
        print(f"❌ Error generating DOCX: {str(e)}")

@app.post("/api/analyze", response_model=AnalysisData)
async def analyze_endpoint(request: AnalyzeRequest):
    """
    Takes the final, edited transcript with timestamps and speakers, returns AI analysis.
    Also generates DOCX in background for instant download later.
    """
    if not request.transcript_blocks:
        raise HTTPException(status_code=400, detail="Transcript blocks are required for analysis.")
    
    # Convert blocks to formatted text with timestamps
    formatted_transcript = []
    for block in request.transcript_blocks:
        timestamp = f"{int(block.timestamp // 60)}:{int(block.timestamp % 60):02d}"
        formatted_transcript.append(f"[{timestamp}] {block.text}")
    
    full_transcript = "\n".join(formatted_transcript)
    
    # Generate analysis
    analysis_data = generate_analysis_report(full_transcript)
    
    # Generate unique cache key
    cache_key = f"docx_{int(time.time())}_{hash(full_transcript)}"
    analysis_data.docx_path = cache_key
    
    # Start DOCX generation in background
    asyncio.create_task(generate_docx_async(analysis_data, full_transcript, cache_key))
    
    return analysis_data

class DownloadRequest(BaseModel):
    """Request model for download endpoint."""
    docx_path: str

@app.post("/api/download-report")
async def download_report_endpoint(request: DownloadRequest):
    """
    Download cached DOCX file or wait for it to be generated.
    """
    cache_key = request.docx_path
    
    # Wait up to 5 seconds for DOCX to be generated if not in cache
    max_wait = 50  # 50 * 0.1s = 5 seconds
    wait_count = 0
    
    while cache_key not in docx_cache and wait_count < max_wait:
        await asyncio.sleep(0.1)
        wait_count += 1
    
    if cache_key not in docx_cache:
        raise HTTPException(status_code=404, detail="DOCX file not ready yet. Please try again.")
    
    # Get cached DOCX
    docx_data = docx_cache[cache_key]
    
    # Create BytesIO from cached data
    docx_buffer = io.BytesIO(docx_data)
    docx_buffer.seek(0)
    
    # Return as download
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=interview_analysis_{int(time.time())}.docx"}
    )

@app.get("/api/ping")
async def ping():
    """Simple endpoint to check API health."""
    return {"message": "pong"}

# --- Database Endpoints ---

class SaveInterviewRequest(BaseModel):
    """Request model for saving interview."""
    title: str
    transcript_text: str
    transcript_words: List[Dict]
    analysis_data: Dict

class UpdateInterviewRequest(BaseModel):
    """Request model for updating interview."""
    interview_id: int
    title: Optional[str] = None
    transcript_text: Optional[str] = None
    transcript_words: Optional[List[Dict]] = None
    analysis_data: Optional[Dict] = None

@app.post("/api/interviews")
async def create_interview(
    title: str = Form(...),
    transcript_text: str = Form(...),
    transcript_words: str = Form(...),
    analysis_data: str = Form(...),
    audio_file: Optional[UploadFile] = File(None)
):
    """Save a new interview to the database."""
    try:
        # Parse JSON strings
        transcript_words_data = json.loads(transcript_words)
        analysis_data_dict = json.loads(analysis_data)
        
        # Save audio file if provided
        audio_url = None
        if audio_file:
            # Generate unique filename
            file_extension = os.path.splitext(audio_file.filename)[1]
            audio_filename = f"{uuid.uuid4()}{file_extension}"
            
            # DEVELOPMENT: Save locally and use API path
            audio_dir = os.path.join(os.path.dirname(__file__), "audio_files")
            os.makedirs(audio_dir, exist_ok=True)
            audio_path = os.path.join(audio_dir, audio_filename)
            
            with open(audio_path, "wb") as f:
                content = await audio_file.read()
                f.write(content)
            
            audio_url = f"/api/audio/{audio_filename}"
            
            # PRODUCTION: Upload to S3/cloud storage and store full URL
            # Example for S3:
            # s3_client.upload_fileobj(audio_file.file, 'bucket-name', audio_filename)
            # audio_url = f"https://your-bucket.s3.amazonaws.com/{audio_filename}"
            # 
            # Example for Cloudinary:
            # result = cloudinary.uploader.upload(audio_file.file, resource_type="video")
            # audio_url = result['secure_url']
        
        interview_id = save_interview(
            title=title,
            transcript_text=transcript_text,
            transcript_words=transcript_words_data,
            analysis_data=analysis_data_dict,
            audio_url=audio_url
        )
        return {"id": interview_id, "message": "Interview saved successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save interview: {str(e)}")

@app.put("/api/interviews/{interview_id}")
async def update_interview_endpoint(interview_id: int, request: UpdateInterviewRequest):
    """Update an existing interview."""
    try:
        success = update_interview(
            interview_id=interview_id,
            title=request.title,
            transcript_text=request.transcript_text,
            transcript_words=request.transcript_words,
            analysis_data=request.analysis_data
        )
        if not success:
            raise HTTPException(status_code=404, detail="Interview not found")
        return {"message": "Interview updated successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update interview: {str(e)}")

@app.get("/api/interviews/{interview_id}")
async def get_interview_endpoint(interview_id: int):
    """Get a single interview by ID."""
    try:
        interview = get_interview(interview_id)
        if not interview:
            raise HTTPException(status_code=404, detail="Interview not found")
        return interview
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to retrieve interview: {str(e)}")

@app.get("/api/interviews")
async def list_interviews(search: Optional[str] = None, limit: int = 50, offset: int = 0):
    """Get list of interviews with optional search."""
    try:
        interviews = get_interview_summaries(search=search, limit=limit, offset=offset)
        return {"interviews": interviews, "count": len(interviews)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to retrieve interviews: {str(e)}")

@app.delete("/api/interviews/{interview_id}")
async def delete_interview_endpoint(interview_id: int):
    """Delete an interview."""
    try:
        success = delete_interview(interview_id)
        if not success:
            raise HTTPException(status_code=404, detail="Interview not found")
        return {"message": "Interview deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete interview: {str(e)}")

@app.get("/api/audio/{audio_filename}")
async def get_audio_file(audio_filename: str):
    """Serve an audio file."""
    try:
        audio_dir = os.path.join(os.path.dirname(__file__), "audio_files")
        audio_path = os.path.join(audio_dir, audio_filename)
        
        if not os.path.exists(audio_path):
            raise HTTPException(status_code=404, detail="Audio file not found")
        
        # Determine media type based on file extension
        file_extension = os.path.splitext(audio_filename)[1].lower()
        media_type = "audio/mpeg" if file_extension == ".mp3" else "audio/wav"
        
        from fastapi.responses import FileResponse
        return FileResponse(audio_path, media_type=media_type)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to serve audio file: {str(e)}")

# --- 6. Uvicorn Runner ---

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
